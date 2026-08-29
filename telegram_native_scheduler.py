#!/usr/bin/env python3
"""
Telegram Quiz + Quote Scheduler
Each session (triggered every 2 hours, or via --loop):
  - Posts N random quiz questions as native Telegram quiz polls, 1-minute delay between each
  - Posts 1 random quote as a plain text message, alongside that same session

Install: pip install python-telegram-bot httpx python-docx
"""

import os
import asyncio
import re
import random
from datetime import datetime
from typing import List, Dict, Optional
import argparse

import httpx
from docx import Document


class TelegramNativeScheduler:
    def __init__(self, bot_token: str):
        """Initialize the scheduler with bot token"""
        self.bot_token = bot_token
        self.base_url = f"https://api.telegram.org/bot{bot_token}"

    # ------------------------------------------------------------------
    # Quiz extraction (unchanged logic)
    # ------------------------------------------------------------------
    async def extract_quizzes_from_docx(self, file_path: str) -> List[Dict]:
        """Extract quizzes from DOCX file"""
        print(f"📂 Reading quizzes: {file_path}")

        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Split by "Question:"
        blocks = [b.strip() for b in re.split(r'Question:', full_text, flags=re.IGNORECASE) if b.strip()]

        print(f"✅ Found {len(blocks)} questions")

        quizzes = []
        skipped = 0

        for idx, block in enumerate(blocks, 1):
            quiz = self._parse_quiz_block(block, idx)

            if quiz:
                quizzes.append(quiz)
            else:
                skipped += 1

        print(f"📋 Valid quizzes: {len(quizzes)}, Skipped: {skipped}\n")
        return quizzes

    def _parse_quiz_block(self, block: str, question_num: int) -> Optional[Dict]:
        """Parse a single quiz block"""

        # Remove explanation (everything after Ans:)
        clean_block = re.split(r'Ans:', block, flags=re.IGNORECASE)[0]

        # Extract question
        question_match = re.match(r'([\s\S]*?)(?=\([a-d]\))', clean_block, re.IGNORECASE)
        if not question_match:
            return None

        question = question_match.group(1).strip()

        # Check for LaTeX
        if re.search(r'\$[\s\S]*?\$|\\[\w\{\}]+', block):
            return None

        # Check for tables
        if re.search(r'^\s*\|[\s\S]*\|', clean_block, re.MULTILINE):
            return None

        # Check for images
        if re.search(r'\[img\]|<img|\.jpg|\.png|\.gif|\.bmp|image:', block, re.IGNORECASE):
            return None

        # Extract options
        option_matches = re.finditer(r'\([a-d]\)\s*([\s\S]*?)(?=\([a-d]\)|$)', clean_block, re.IGNORECASE)
        options = []
        for match in option_matches:
            opt = match.group(1).strip()
            if opt:
                options.append(opt)

        # Extract answer
        answer_match = re.search(r'Ans:\s*([a-d])', block, re.IGNORECASE)
        if not answer_match:
            return None

        correct_idx = ord(answer_match.group(1).lower()) - ord('a')

        # Validation
        if len(options) < 2 or len(options) > 10:
            return None

        if len(question) > 300:
            return None

        if correct_idx >= len(options):
            return None

        if any(len(opt) > 100 for opt in options):
            return None

        return {
            'number': question_num,
            'question': question,
            'options': options,
            'correct_option_id': correct_idx
        }

    # ------------------------------------------------------------------
    # Quote extraction
    # ------------------------------------------------------------------
    async def extract_quotes_from_docx(self, file_path: str) -> List[str]:
        """Extract quotes from DOCX file.

        Expected format: one quote per paragraph, each starting with
        'Quote:' followed by the quote text (may contain a line break).
        """
        print(f"📂 Reading quotes: {file_path}")

        doc = Document(file_path)
        quotes = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Strip a leading "Quote:" label if present
            match = re.match(r'^Quote:\s*(.+)$', text, re.IGNORECASE | re.DOTALL)
            quote_text = match.group(1).strip() if match else text

            if quote_text:
                quotes.append(quote_text)

        print(f"✅ Found {len(quotes)} quotes\n")
        return quotes

    # ------------------------------------------------------------------
    # Sending
    # ------------------------------------------------------------------
    async def send_poll(self, client: httpx.AsyncClient, chat_id: int, quiz: Dict) -> bool:
        """Send a single quiz poll to a chat"""
        try:
            payload = {
                'chat_id': chat_id,
                'question': quiz['question'],
                'options': quiz['options'],
                'type': 'quiz',
                'correct_option_id': quiz['correct_option_id'],
                # Telegram does NOT allow non-anonymous polls in channels
                # (only in groups/supergroups). Channels require
                # is_anonymous=True or sendPoll fails outright.
                'is_anonymous': True,
            }

            response = await client.post(
                f"{self.base_url}/sendPoll",
                json=payload,
                timeout=30.0
            )

            result = response.json()
            ok = result.get('ok', False)
            if not ok:
                # Surface the real Telegram error instead of failing silently
                print(f"    ↳ Telegram error: {result.get('description', result)}")
            return ok

        except Exception as e:
            print(f"Error sending poll: {str(e)}")
            return False

    async def send_message(self, client: httpx.AsyncClient, chat_id: int, text: str) -> bool:
        """Send a plain text message (used for quotes) to a chat"""
        try:
            payload = {
                'chat_id': chat_id,
                'text': text,
            }

            response = await client.post(
                f"{self.base_url}/sendMessage",
                json=payload,
                timeout=30.0
            )

            result = response.json()
            ok = result.get('ok', False)
            if not ok:
                print(f"    ↳ Telegram error: {result.get('description', result)}")
            return ok

        except Exception as e:
            print(f"Error sending message: {str(e)}")
            return False

    # ------------------------------------------------------------------
    # Session
    # ------------------------------------------------------------------
    async def post_session(
        self,
        quizzes: List[Dict],
        quotes: List[str],
        chat_ids: List[int],
        quote_chat_ids: List[int],
        num_questions: int,
        random_seed: Optional[int] = None,
    ):
        """Post one quote (random) plus N random quiz questions (1-minute
        delay between each quiz poll) to the given chats.
        """

        if random_seed is not None:
            random.seed(random_seed)

        selected_quizzes = random.sample(quizzes, min(num_questions, len(quizzes)))
        selected_quizzes = sorted(selected_quizzes, key=lambda x: x['number'])

        selected_quote = random.choice(quotes) if quotes else None

        print(f"🎲 RANDOM SELECTION: {len(selected_quizzes)} questions")
        print(f"📌 Questions: {[q['number'] for q in selected_quizzes]}")
        print(f"⏱️  Delay: 1 minute between each quiz poll")
        print(f"👥 Quiz groups: {len(chat_ids)}")
        if selected_quote:
            print(f"💬 Quote groups: {len(quote_chat_ids)}")
        print("─" * 60)

        async with httpx.AsyncClient() as client:
            # Post the quote first, once per session
            if selected_quote:
                print(f"\n💬 Sending quote...")
                for chat_id in quote_chat_ids:
                    success = await self.send_message(client, chat_id, selected_quote)
                    status = "✅ Sent" if success else "❌ Failed"
                    print(f"  {status} to chat {chat_id}")

            for idx, quiz in enumerate(selected_quizzes):
                print(f"\n[{idx+1}/{len(selected_quizzes)}] Sending Q{quiz['number']}...")

                for chat_id in chat_ids:
                    success = await self.send_poll(client, chat_id, quiz)
                    status = "✅ Sent" if success else "❌ Failed"
                    print(f"  {status} to chat {chat_id}")

                if idx < len(selected_quizzes) - 1:
                    print(f"  ⏳ Waiting 60 seconds before next poll...")
                    for remaining in range(60, 0, -10):
                        await asyncio.sleep(10)
                        if remaining > 10:
                            print(f"     {remaining-10}s remaining...")

        print("\n" + "─" * 60)
        print("✅ SESSION COMPLETE")
        print(f"⏰ Total time: ~{max(len(selected_quizzes)-1, 0)} minutes")
        print("📱 Check your Telegram groups\n")


async def main():
    """Main function - posts quiz(zes) + a quote, every 2 hours if --loop"""
    parser = argparse.ArgumentParser(
        description='Post random Telegram quizzes plus a quote, each session',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Post quizzes + a quote once
  python telegram_native_scheduler.py --token YOUR_BOT_TOKEN --file quiz.docx --quotes-file quotes.docx --chat 123456789

  # Send the quote to a different chat than the quizzes
  python telegram_native_scheduler.py --token YOUR_BOT_TOKEN --file quiz.docx --quotes-file quotes.docx --chat 123456789 --quote-chat -1003997553872

  # With reproducible seed
  python telegram_native_scheduler.py --token YOUR_BOT_TOKEN --file quiz.docx --quotes-file quotes.docx --chat 123456789 --seed 42

  # Keep running every 2 hours
  python telegram_native_scheduler.py --token YOUR_BOT_TOKEN --file quiz.docx --quotes-file quotes.docx --chat 123456789 --loop
        """
    )

    parser.add_argument('--token', required=True, help='Telegram Bot Token')
    parser.add_argument('--file', required=True, help='DOCX file with quizzes')
    parser.add_argument('--quotes-file', help='DOCX file with quotes (one per session, random)')
    parser.add_argument('--chat', type=int, action='append', required=True, dest='chats', help='Chat ID for quizzes')
    parser.add_argument('--quote-chat', type=int, action='append', dest='quote_chats',
                         help='Chat ID for quotes (defaults to the same as --chat if omitted)')
    parser.add_argument('--num-questions', type=int, default=30, help='Questions per session (default: 30)')
    parser.add_argument('--seed', type=int, help='Random seed (optional)')
    parser.add_argument('--loop', action='store_true', help='Post every 2 hours indefinitely')

    args = parser.parse_args()

    if not os.path.exists(args.file):
        print(f"❌ File not found: {args.file}")
        return

    if args.quotes_file and not os.path.exists(args.quotes_file):
        print(f"❌ Quotes file not found: {args.quotes_file}")
        return

    scheduler = TelegramNativeScheduler(args.token)

    quizzes = await scheduler.extract_quizzes_from_docx(args.file)
    quotes = await scheduler.extract_quotes_from_docx(args.quotes_file) if args.quotes_file else []

    if not quizzes:
        print("❌ No valid quizzes found")
        return

    if len(quizzes) < args.num_questions:
        print(f"⚠️  Warning: Only {len(quizzes)} questions available (need {args.num_questions})")

    quote_chats = args.quote_chats if args.quote_chats else args.chats

    if args.loop:
        print(f"🔄 Running every 2 hours (press Ctrl+C to stop)\n")
        counter = 1
        while True:
            print(f"\n{'='*60}")
            print(f"📅 Session {counter} - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"{'='*60}")

            await scheduler.post_session(
                quizzes, quotes, args.chats, quote_chats, args.num_questions, args.seed
            )

            print(f"⏰ Next session in 2 hours...")
            counter += 1
            await asyncio.sleep(7200)
    else:
        await scheduler.post_session(
            quizzes, quotes, args.chats, quote_chats, args.num_questions, args.seed
        )


if __name__ == '__main__':
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n\n⚠️  Stopped by user")
